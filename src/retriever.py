import os
from docx import Document as DocxDocument

CONTRACTS_DIR = "../data/contracts"

def read_docx(file_path):
    doc = DocxDocument(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

def load_contracts():
    contracts = []
    for filename in os.listdir(CONTRACTS_DIR):
        if filename.endswith(".docx"):
            file_path = os.path.join(CONTRACTS_DIR, filename)
            contracts.append(read_docx(file_path))
    return contracts

def read_docx_from_file(file):
    doc = DocxDocument(file)
    return "\n".join([para.text for para in doc.paragraphs])

def load_contracts_from_uploaded_files(uploaded_files):
    contracts = [read_docx_from_file(file) for file in uploaded_files]
    return contracts
